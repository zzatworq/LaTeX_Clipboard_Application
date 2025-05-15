import base64
import html
import tempfile
import time
import tkinter as tk
from tkinter import ttk, messagebox, filedialog
import threading
import sys
import logging
import os
import json
from .components import create_settings_frame, create_actions_frame, create_io_frame
from src.utils.clipboard import set_clipboard_html, get_clipboard_text
from src.utils.latex import check_latex, find_latex_equations
from src.utils.image import render_latex_to_image, is_image_empty, image_to_bytes
from src.config.settings import configure_logging

class LatexClipboardApp:
    def __init__(self, root):
        self.root = root
        self.root.title("LaTeX Clipboard Monitor")
        self.monitoring = False
        self.monitor_thread = None
        self.stop_event = threading.Event()
        self.last_images = []
        self.last_text = ""
        self.last_equations = None
        self.defaults_file = os.path.join("configs", "defaults.json")
        self.logger_enabled = tk.BooleanVar(value=True)

        if not check_latex():
            messagebox.showerror("LaTeX Not Found", "LaTeX distribution (e.g., MiKTeX) with latex and dvipng required.")
            sys.exit(1)

        configure_logging(self.logger_enabled.get())
        self.load_defaults()
        self.root.state('normal')
        self.root.attributes('-topmost', True)
        self.root.update()
        self.root.attributes('-topmost', False)
        self.root.focus_force()

        self.create_gui()
        logging.info("Application initialized.")

    def load_defaults(self):
        defaults = {
            "mode": "Matplotlib",
            "text_color": "white",
            "font_size": "12",
            "dpi": "300",
            "only_images": False,
            "logger_enabled": True
        }
        try:
            if os.path.exists(self.defaults_file):
                with open(self.defaults_file, 'r') as f:
                    loaded = json.load(f)
                defaults.update({k: v for k, v in loaded.items() if k in defaults})
                logging.info(f"Loaded defaults: {defaults}")
            else:
                logging.info("Using fallback defaults")
        except Exception as e:
            logging.error(f"Failed to load defaults: {e}")
        self.default_settings = defaults
        self.logger_enabled.set(defaults["logger_enabled"])

    def save_defaults(self, settings):
        try:
            os.makedirs(os.path.dirname(self.defaults_file), exist_ok=True)
            with open(self.defaults_file, 'w') as f:
                json.dump(settings, f, indent=4)
            logging.info(f"Saved defaults: {settings}")
        except Exception as e:
            logging.error(f"Failed to save defaults: {e}")
            messagebox.showerror("Save Defaults Failed", f"Error: {e}")

    def create_gui(self):
        main_frame = ttk.Frame(self.root, padding="10")
        main_frame.grid(row=0, column=0, sticky="nsew")
        self.root.columnconfigure(0, weight=1)
        self.root.rowconfigure(0, weight=1)

        self.settings_frame = create_settings_frame(main_frame, self.default_settings, self.logger_enabled, self.validate_inputs)
        self.actions_frame = create_actions_frame(main_frame, self.toggle_monitoring, self.test_render, self.save_as_docx, self.open_defaults_dialog)
        self.io_frame, self.text_input, self.status_var = create_io_frame(main_frame, self.render_input_text)

        self.root.protocol("WM_DELETE_WINDOW", self.on_closing)

    def validate_inputs(self, font_size_var, dpi_var):
        try:
            font_size = int(font_size_var.get())
            dpi = int(dpi_var.get())
            if not (10 <= font_size <= 50):
                raise ValueError("Font size must be 10-50")
            if not (100 <= dpi <= 600):
                raise ValueError("DPI must be 100-600")
            return True
        except ValueError as e:
            logging.error(f"Input validation failed: {e}")
            messagebox.showerror("Invalid Input", str(e))
            return False

    def toggle_monitoring(self):
        configure_logging(self.logger_enabled.get())
        if not self.monitoring:
            if not self.validate_inputs(self.settings_frame.font_size_var, self.settings_frame.dpi_var):
                return
            self.monitoring = True
            self.status_var.set("Monitoring")
            self.actions_frame.toggle_button.configure(text="Stop Monitoring")
            self.disable_gui()
            self.stop_event.clear()
            self.root.after(1000, self.start_monitor_thread)
            logging.info("Started clipboard monitoring")
        else:
            self.monitoring = False
            self.status_var.set("Stopped")
            self.actions_frame.toggle_button.configure(text="Start Monitoring")
            self.enable_gui()
            self.stop_event.set()
            logging.info("Stopped clipboard monitoring")

    def disable_gui(self):
        for widget in [self.settings_frame.mode_menu, self.settings_frame.color_menu, self.settings_frame.font_size_spin,
                       self.settings_frame.dpi_spin, self.actions_frame.test_button, self.actions_frame.save_button,
                       self.actions_frame.defaults_button, self.io_frame.render_button, self.text_input]:
            widget.configure(state="disabled")

    def enable_gui(self):
        for widget in [self.settings_frame.mode_menu, self.settings_frame.color_menu, self.settings_frame.font_size_spin,
                       self.settings_frame.dpi_spin, self.actions_frame.test_button, self.actions_frame.save_button,
                       self.actions_frame.defaults_button, self.io_frame.render_button, self.text_input]:
            widget.configure(state="normal")

    def start_monitor_thread(self):
        self.monitor_thread = threading.Thread(target=self.monitor_clipboard, daemon=True)
        self.monitor_thread.start()

    def render_input_text(self):
        configure_logging(self.logger_enabled.get())
        if not self.validate_inputs(self.settings_frame.font_size_var, self.settings_frame.dpi_var):
            return
        text = self.text_input.get("1.0", tk.END).strip()
        if not text:
            messagebox.showwarning("No Input", "Please enter text to render.")
            return
        self.process_text(text, "input")

    def test_render(self):
        configure_logging(self.logger_enabled.get())
        if not self.validate_inputs(self.settings_frame.font_size_var, self.settings_frame.dpi_var):
            return
        from templates.test_string import TEST_STRING
        self.process_text(TEST_STRING, "test")

    def process_text(self, text, mode):
        logging.info(f"Rendering {mode} text: {text[:100]}...")
        try:
            equations = find_latex_equations(text)
            images = [
                render_latex_to_image(eq, self.settings_frame.color_var.get(), int(self.settings_frame.font_size_var.get()),
                                      int(self.settings_frame.dpi_var.get()), mode=self.settings_frame.mode_var.get())
                for eq in equations['equations']
            ]
            images = [img for img in images if img]
            if images:
                self.copy_images(images, mode == "test", text, equations)
                self.status_var.set(f"Copied {len(images)} images")
                messagebox.showinfo(f"{mode.capitalize()} Render", f"Copied {len(images)} images")
            else:
                self.status_var.set("No valid images")
                messagebox.showerror(f"{mode.capitalize()} Render", "Failed to render images")
        except Exception as e:
            logging.error(f"{mode.capitalize()} render failed: {e}")
            messagebox.showerror(f"{mode.capitalize()} Render Failed", f"Error: {e}")
            self.status_var.set(f"{mode.capitalize()} render failed")

    def save_as_docx(self):
        configure_logging(self.logger_enabled.get())
        if not self.last_images:
            messagebox.showwarning("No Images", "No images available to save.")
            return
        from docx import Document
        from docx.shared import Pt
        import filedialog
        file_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word Documents", "*.docx")])
        if not file_path:
            return
        try:
            doc = Document()
            only_images = self.settings_frame.only_images_var.get()
            if only_images or not self.last_text or not self.last_equations['matches']:
                for img in self.last_images:
                    with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as tmp:
                        img.save(tmp.name, format='PNG')
                        doc.add_picture(tmp.name, width=Pt(300))
                        os.unlink(tmp.name)
            else:
                last_pos = 0
                img_index = 0
                font_size = int(self.settings_frame.font_size_var.get())
                for match in self.last_equations['matches']:
                    start, end = match['start'], match['end']
                    text_segment = self.last_text[last_pos:start].strip()
                    if text_segment:
                        doc.add_paragraph(text_segment).runs[0].font.size = Pt(font_size)
                    if img_index < len(self.last_images):
                        with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as tmp:
                            self.last_images[img_index].save(tmp.name, format='PNG')
                            doc.add_paragraph().add_run().add_picture(tmp.name, width=Pt(300))
                            os.unlink(tmp.name)
                        img_index += 1
                    last_pos = end
                if remaining := self.last_text[last_pos:].strip():
                    doc.add_paragraph(remaining).runs[0].font.size = Pt(font_size)
            doc.save(file_path)
            logging.info(f"Saved DOCX to {file_path}")
            messagebox.showinfo("Save Successful", f"Saved to {file_path}")
            os.startfile(file_path)
        except Exception as e:
            logging.error(f"Failed to save DOCX: {e}")
            messagebox.showerror("Save Failed", f"Error: {e}")

    def open_defaults_dialog(self):
        configure_logging(self.logger_enabled.get())
        dialog = tk.Toplevel(self.root)
        dialog.title("Set Default Settings")
        dialog.transient(self.root)
        dialog.grab_set()
        dialog.resizable(False, False)

        frame = ttk.Frame(dialog, padding="15")
        frame.grid(row=0, column=0, sticky="nsew")

        vars = {
            "mode": tk.StringVar(value=self.default_settings["mode"]),
            "text_color": tk.StringVar(value=self.default_settings["text_color"]),
            "font_size": tk.StringVar(value=self.default_settings["font_size"]),
            "dpi": tk.StringVar(value=self.default_settings["dpi"]),
            "only_images": tk.BooleanVar(value=self.default_settings["only_images"]),
            "logger_enabled": tk.BooleanVar(value=self.default_settings["logger_enabled"])
        }

        ttk.Label(frame, text="Default Render Mode:").grid(row=0, column=0, padx=10, pady=10, sticky="e")
        ttk.OptionMenu(frame, vars["mode"], self.default_settings["mode"], "Matplotlib", "Standalone").grid(row=0, column=1, padx=10, pady=10, sticky="w")

        ttk.Label(frame, text="Default Text Color:").grid(row=1, column=0, padx=10, pady=10, sticky="e")
        ttk.OptionMenu(frame, vars["text_color"], self.default_settings["text_color"], "white", "black", "red", "blue", "green").grid(row=1, column=1, padx=10, pady=10, sticky="w")

        ttk.Label(frame, text="Default Font Size:").grid(row=2, column=0, padx=10, pady=10, sticky="e")
        ttk.Spinbox(frame, from_=10, to=50, width=10, textvariable=vars["font_size"]).grid(row=2, column=1, padx=10, pady=10, sticky="w")

        ttk.Label(frame, text="Default DPI:").grid(row=3, column=0, padx=10, pady=10, sticky="e")
        ttk.Spinbox(frame, from_=100, to=600, width=10, textvariable=vars["dpi"]).grid(row=3, column=1, padx=10, pady=10, sticky="w")

        ttk.Checkbutton(frame, text="Default Only Images", variable=vars["only_images"]).grid(row=4, column=0, columnspan=2, padx=10, pady=10, sticky="w")
        ttk.Checkbutton(frame, text="Enable Logging", variable=vars["logger_enabled"]).grid(row=5, column=0, columnspan=2, padx=10, pady=10, sticky="w")

        def save():
            try:
                font_size = int(vars["font_size"].get())
                dpi = int(vars["dpi"].get())
                if not (10 <= font_size <= 50 and 100 <= dpi <= 600):
                    raise ValueError("Font size 10-50, DPI 100-600")
                new_defaults = {k: v.get() for k, v in vars.items()}
                self.default_settings = new_defaults
                self.save_defaults(new_defaults)
                self.settings_frame.mode_var.set(new_defaults["mode"])
                self.settings_frame.color_var.set(new_defaults["text_color"])
                self.settings_frame.font_size_var.set(new_defaults["font_size"])
                self.settings_frame.dpi_var.set(new_defaults["dpi"])
                self.settings_frame.only_images_var.set(new_defaults["only_images"])
                self.logger_enabled.set(new_defaults["logger_enabled"])
                messagebox.showinfo("Defaults Saved", "Default settings updated.")
                dialog.destroy()
            except ValueError as e:
                messagebox.showerror("Invalid Input", str(e))

        ttk.Button(frame, text="Save", command=save).grid(row=6, column=0, padx=10, pady=15)
        ttk.Button(frame, text="Cancel", command=dialog.destroy).grid(row=6, column=1, padx=10, pady=15)

        dialog.update_idletasks()
        width = frame.winfo_reqwidth() + 20
        height = frame.winfo_reqheight() + 20
        x = self.root.winfo_rootx() + (self.root.winfo_width() - width) // 2
        y = self.root.winfo_rooty() + (self.root.winfo_height() - height) // 2
        dialog.geometry(f"{width}x{height}+{x}+{y}")

    def copy_images(self, images, test_mode, original_text, equations):
        configure_logging(self.logger_enabled.get())
        if not images:
            logging.info("No images to copy")
            return
        self.last_images = [img for img in images if img and not is_image_empty(img)]
        self.last_text = original_text
        self.last_equations = equations

        html_content = (
            "<style>"
            f"body {{color: {self.settings_frame.color_var.get()}; font-family: Arial, sans-serif; font-size: {self.settings_frame.font_size_var.get()}pt; line-height: 1.5;}}"
            "p, div, span {color: inherit !important;}"
            "img {vertical-align: middle; margin: 2px 0;}"
            "</style>"
        )

        if test_mode or (original_text and equations['matches']):
            if self.settings_frame.only_images_var.get():
                html_content += "".join(
                    f'<img src="data:image/png;base64,{base64.b64encode(image_to_bytes(img)).decode()}" style="vertical-align: middle; margin: 2px 0;">'
                    for img in self.last_images
                )
            else:
                last_pos = 0
                img_index = 0
                for match in equations['matches']:
                    start, end = match['start'], match['end']
                    text_segment = html.escape(original_text[last_pos:start]).replace('\n', '<br>')
                    html_content += f'<span>{text_segment}</span>'
                    if img_index < len(self.last_images):
                        html_content += f'<img src="data:image/png;base64,{base64.b64encode(image_to_bytes(self.last_images[img_index])).decode()}" style="vertical-align: middle; margin: 2px 0;">'
                        img_index += 1
                    last_pos = end
                html_content += f'<span>{html.escape(original_text[last_pos:]).replace("\n", "<br>")}</span>'
        else:
            html_content += "<br>".join(
                f'<img src="data:image/png;base64,{base64.b64encode(image_to_bytes(img)).decode()}" style="vertical-align: middle; margin: 2px 0;">'
                for img in self.last_images
            )

        try:
            set_clipboard_html(html_content)
            self.status_var.set(f"Copied {len(self.last_images)} images")
            logging.info(f"Copied {len(self.last_images)} images")
        except Exception as e:
            logging.error(f"Failed to copy images: {e}")
            self.status_var.set("Error copying images")

    def monitor_clipboard(self):
        configure_logging(self.logger_enabled.get())
        last_sequence = None
        while not self.stop_event.is_set():
            try:
                import win32clipboard
                current_sequence = win32clipboard.GetClipboardSequenceNumber()
                if current_sequence != last_sequence:
                    last_sequence = current_sequence
                    text = get_clipboard_text()
                    if text:
                        logging.info(f"New clipboard content: {text[:100]}...")
                        equations = find_latex_equations(text)
                        if equations['equations']:
                            images = [
                                render_latex_to_image(eq, self.settings_frame.color_var.get(), int(self.settings_frame.font_size_var.get()),
                                                      int(self.settings_frame.dpi_var.get()), mode=self.settings_frame.mode_var.get())
                                for eq in equations['equations']
                            ]
                            images = [img for img in images if img]
                            if images:
                                self.copy_images(images, False, text, equations)
                            else:
                                self.status_var.set("No valid images")
                        else:
                            self.status_var.set("No equations found")
                time.sleep(1)
            except Exception as e:
                logging.error(f"Clipboard monitoring error: {e}")
                self.status_var.set("Monitoring error")
                time.sleep(1)

    def on_closing(self):
        if self.monitoring:
            self.stop_event.set()
            if self.monitor_thread:
                self.monitor_thread.join(timeout=1.0)
        self.root.destroy()
        logging.info("Application closed")